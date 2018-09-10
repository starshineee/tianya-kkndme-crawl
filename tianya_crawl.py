import requests,time
from docx import Document
from bs4 import BeautifulSoup

base_url='http://bbs.tianya.cn/m/'
headers = {
        'User-Agent':'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_12_3) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.36'
    }
layer_ind = 0
page = 'post_author-house-252774-1.shtml'

cookie = {yourcookie} # 自己登录天涯并从浏览器中得到 yourcookie
cookies = {}
for line in cookie.split(';'):
    name,value=line.strip().split('=',1)
    cookies[name]=value


document = Document()
while True:
    print(page)
    r=requests.get(url=base_url+page,headers=headers,cookies=cookies)
    soup=BeautifulSoup(r.text,'html.parser')
    contents=soup.find_all('div',class_='bd')
    for content in contents:
        paragraphs=content.find_all("p",recursive=False)
        document.add_heading('层{}'.format(layer_ind),level=2)
        for p in paragraphs:
            document.add_paragraph(p.get_text())
        document.add_paragraph('==============================')
        layer_ind+=1

    page=soup.find('a',class_='u-btn next-btn')
    if page==None:
        break
    else:
        page=page.get('href')
        time.sleep(1)
document.save('demo.docx')